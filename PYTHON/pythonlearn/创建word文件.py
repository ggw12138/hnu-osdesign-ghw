# -*- encoding:UTF-8 -*-
# 誉天暑期训练营 python-docx
# 数据写入
from docx import Document
from docx.shared import Inches
# 创建对象
document = Document()
document.add_heading('Python暑期训练营', 0)  # 0代表的是标题的类型，共有4中类型
# 添加正文内容并设置部分内容格式
p = document.add_paragraph('Python 爬虫开发---')
# 设置内容加粗
p.runs[0].bold = True
# 添加内容并加粗
p.add_run('数据存储-').bold = True
# 添加内容
p.add_run('Word-')
# 设置字体
p.add_run('还给老师了-').italic = True
# 添加正文，设置"项目符号"
document.add_paragraph(
    '项目符号1', style='ListBullet'
)
document.add_paragraph(
    '项目符号2', style='ListNumber'
)
# 添加图片
document.add_picture('ball.png', width=Inches(1.25))
# 添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for i in range(2):
    row_cells = table.add_row().cells
    row_cells[0].text = 'a'
    row_cells[1].text = 'b'
    row_cells[2].text = 'c'
# 保存文件
document.add_page_break()
document.save('test.docx')

'''
在word中写入数据的整体思路
1、创建生成临时的word对象
2、分别使用add_paragraph()和add_heading()对word对象添加标题和正文内容
3、如何像设置正文内容字体加粗或斜体，可以将正文内容p对象的属性p.run[0].bold=True和add_run('Xx').italic= True设置为斜体
4、如果要插入图片和添加表格，可以在word
'''
