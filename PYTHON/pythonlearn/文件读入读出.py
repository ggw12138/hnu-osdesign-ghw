from docx.shared import Inches
from docx import Document
import xlrd
import xlwt

# 新建一个excel文件
wb = xlwt.Workbook()
# 新建一个sheet
ws = wb.add_sheet('python1', cell_overwrite_ok=True)
# 定义字体对齐方式对象
alignment = xlwt.Alignment()
# 设置水平方向  中心对齐
alignment.horz = xlwt.Alignment.HORZ_CENTER
# 设置垂直方向  中心对齐
alignment.vert = xlwt.Alignment.VERT_CENTER
# 定义格式化对象
style = xlwt.XFStyle()
style.alignment = alignment
# 合并单元格
# ws.write_merge(开始行,结束行,开始列,结束列,内容,格式)
ws.write_merge(0, 0, 0, 5, 'pythonlearn', style)
# 写入数据
for i in range(2, 7):
    for k in range(5):
        ws.write(i, k, i+k)
        # ws.write(行,列,内容)

# 保存文件
wb.save('file.xls')

# excel内容的读取
wb = xlrd.open_workbook('file.xls')
# 获取sheets总数
ws_count = wb.nsheets
print('Sheets总数:', ws_count)
# 通过索引顺序获取Sheets
# ws=wb.sheets()[0]
ws = wb.sheet_by_name('Python')
# 获取整行内容
row_value = ws.row_values(3)
print('第4行数据:', row_value)
# 获取整列的值
row_col = ws.col_values(3)
print('第4列数据', row_col)
# 获取总行数和总列数
nrows = ws.nrows
ncols = ws.ncols
print('总行数:', nrows, '总列数:', ncols)
cells = ws.cell(2, 3).value
print(cells)

'''
# 新建word文档
# 创建对象
document = Document()
document.add_heading('Python暑期训练营', 0)  # 0代表的是标题的类型，共有4中类型
# 添加正文内容并设置部分内容格式
p = document.add_paragraph('Python 爬虫开发---')
# 设置内容加粗
p.runs[0].bold = True
# 添加内容并加粗
p.add_run('数据存储-').bold = True
# 添加内容
p.add_run('Word-')
# 设置字体  斜体
p.add_run('还给老师了-').italic = True
# 添加正文，设置"项目符号"
document.add_paragraph(
    '项目符号1', style='ListBullet'
)
document.add_paragraph(
    '项目符号2', style='ListNumber'
)
# 添加图片
document.add_picture('D:\VSCode\code\PYTHON\pythonlearn\ball.png', width=Inches(1.25))
# 添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for i in range(2):
    row_cells = table.add_row().cells
    row_cells[0].text = 'a'
    row_cells[1].text = 'b'
    row_cells[2].text = 'c'
# 保存文件
document.add_page_break()
document.save('test.docx')
'''